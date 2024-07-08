import os
import sys
from typing import List

from PyPDF2 import PdfReader, PdfWriter


def save_to_word(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save to word data
    :param path: path to save to file
    :param result_data: Data to save to word file
    :param data: Data to save to word file
    :param textboxes: Text boxes to save to word file
    :return: True on success, False on failure + Exception
    """
    try:
        from docx import Document
    except ImportError:
        return False, Exception('docx module is not installed')

    if data == []:
        return False, Exception('Data cannot be empty')

    data_ptr = data[0].get('mdth', None)
    if not data_ptr:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    document = Document()
    document.add_heading("РЕЗУЛЬТАТЫ", level=1)

    table_main = document.add_table(rows=1, cols=3)

    for i, item in enumerate(data_ptr):
        row_cells_main = table_main.add_row().cells
        row_cells_main[0].text = item.get("idx", "")
        row_cells_main[1].text = item.get("name", "")

        if i < len(textboxes):
            row_cells_main[2].text = textboxes[i].text()

    if result_data:
        document.add_paragraph("\n")
        document.add_heading("Таблица результатов", level=2)

        table_results = document.add_table(rows=len(result_data), cols=len(result_data[0]) if result_data else 0)
        for row, row_data in enumerate(result_data):
            for col, value in enumerate(row_data):
                cell = table_results.cell(row, col)
                cell.text = str(value)
    try:
        document.save(f'{path}/export/Untitled.docx')
    except Exception as e:
        return False, Exception(e)
    else:
        return True, None


def append_to_word(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
       Save to word data
       :param path: path to save to file
       :param result_data: Data to save to word file
       :param data: Data to save to word file
       :param textboxes: Text boxes to save to word file
       :return: True on success, False on failure + Exception
       """
    try:
        from docx import Document
    except ImportError:
        return False, Exception('docx module is not installed')

    try:
        if not os.path.exists(path):
            os.makedirs(path)

        if not os.path.exists(f'{path}/export/Untitled.docx'):
            return save_to_word(path, data, textboxes, result_data)

        file_path = os.path.join(path, 'export', 'Untitled.docx')


        if not os.path.exists(file_path):
            document = Document()
        else:
            document = Document(file_path)

        if data == []:
            return False, Exception('Data cannot be empty')

        data_ptr = data[0].get('mdth', None)
        if not data_ptr:
            return False, Exception('Data cannot be empty')

        textboxes_ptr: List = textboxes
        if not textboxes_ptr:
            return False, Exception('Text boxes cannot be empty')


        document.add_heading("РЕЗУЛЬТАТЫ", level=1)

        table_main = document.add_table(rows=1, cols=3)
        hdr_cells = table_main.rows[0].cells
        hdr_cells[0].text = '№ п/п'
        hdr_cells[1].text = 'Показатель'
        hdr_cells[2].text = 'Значение'

        for i, item in enumerate(data_ptr):
            row_cells_main = table_main.add_row().cells
            row_cells_main[0].text = item.get("idx", "")
            row_cells_main[1].text = item.get("name", "")

            if i < len(textboxes):
                row_cells_main[2].text = textboxes[i].text()

        if result_data:
            document.add_paragraph("\n")
            document.add_heading("Таблица результатов", level=2)

            table_results = document.add_table(rows=1, cols=len(result_data[0]))
            hdr_cells = table_results.rows[0].cells
            for col, value in enumerate(result_data[0]):
                hdr_cells[col].text = f"Column {col + 1}"

            for row_data in result_data:
                row_cells_results = table_results.add_row().cells
                for col, value in enumerate(row_data):
                    row_cells_results[col].text = str(value)

        document.save(file_path)
        return True, None
    except Exception as e:
        return False, Exception(e)


def save_to_excel(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to excel file
    :param path: path to save to file
    :param data: Data to save to excel file
    :param textboxes: Text boxes to save to excel file
    :return: True on success, False on failure + Exception
    """
    try:
        import pandas as pd
    except ImportError:
        return False, Exception('Pandas is not installed')

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    data_list: List = [{'№ п/п': item.get("idx", ""), 'Показатель': item.get("name", ""),
                        'data': textboxes_ptr[i].text()} for i, item in enumerate(data_prt)]
    try:
        df_data = pd.DataFrame(data_list)
        df_result = pd.DataFrame(result_data) if result_data else None
    except Exception as e:
        return False, Exception(e)
    else:
        try:
            with pd.ExcelWriter(f'{path}/export/Untitled.xlsx') as writer:
                df_data.to_excel(writer, sheet_name='Data', index=False)
                if df_result is not None:
                    df_result.to_excel(writer, sheet_name='Result Data', index=False)
        except Exception as e:
            return False, Exception(e)
        else:
            return True, None


def append_to_excel(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to excel file
    :param path: path to save to file
    :param data: Data to save to excel file
    :param textboxes: Text boxes to save to excel file
    :param result_data: Additional result data to save
    :return: True on success, False on failure + Exception
    """
    try:
        import pandas as pd
    except ImportError:
        return False, Exception('Pandas is not installed')

    try:
        if data == []:
            return False, Exception('Data cannot be empty')

        data_prt = data[0].get('mdth', None)
        if not data_prt:
            return False, Exception('Data cannot be empty')

        textboxes_ptr: List = textboxes
        if not textboxes_ptr:
            return False, Exception('Text boxes cannot be empty')

        data_list: List = [{'data': item.text()} for item in textboxes_ptr]

        df_data = pd.DataFrame(data_list)
        append_result = []
        if result_data:
            for item in result_data:
                append_result.append(item[2])
            df_result = pd.DataFrame(append_result) if append_result else pd.DataFrame()
        else:
            df_result = pd.DataFrame()
            
        if not os.path.exists(f'{path}/export/Untitled.xlsx'):
            return save_to_excel(path, data, textboxes, result_data)
        
        try:
            existing_df_data = pd.read_excel(f'{path}/export/Untitled.xlsx', sheet_name='Data')
            existing_df_result = pd.read_excel(f'{path}/export/Untitled.xlsx', sheet_name='Result Data')
        except FileNotFoundError:
            existing_df_data = pd.DataFrame()
            existing_df_result = pd.DataFrame()
            
        combined_df_data = pd.concat([existing_df_data, df_data], axis=1)
        combined_df_result = pd.concat([existing_df_result, df_result], axis=1)

        with pd.ExcelWriter(f'{path}/export/Untitled.xlsx', engine='openpyxl', mode='a',
                            if_sheet_exists='replace') as writer:
            combined_df_data.to_excel(writer, sheet_name='Data', index=False)
            combined_df_result.to_excel(writer, sheet_name='Result Data', index=False)

        return True, None
    except Exception as e:
        return False, Exception(e)


def save_to_pdf(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to pdf file
    :param path: path to save to file
    :param data: Data to save to pdf file
    :param textboxes: Text boxes to save to pdf file
    :param result_data: Result data to save to pdf file
    :return: True on success, False on failure + Exception
    """
    try:
        from fpdf import FPDF
    except ImportError:
        return False, Exception('FPDF is not installed')

    if not data:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    if not textboxes:
        return False, Exception('Text boxes cannot be empty')

    class PDF(FPDF):
        """
        PDF class for saving data to pdf file
        """

        def header(self) -> None:
            """
            Set header of pdf file
            :return: None
            """
            self.set_font('Arial', 'B', 12)
            self.cell(0, 10, 'Data from Textboxes', 0, 1, 'C')

    pdf = PDF()
    pdf.add_page()

    if sys.platform.startswith('win32'):
        arial_path = 'C:\\Windows\\Fonts\\arial.ttf'
        arial_bold_path = 'C:\\Windows\\Fonts\\arialbd.ttf'

        # Check if font files exist
        if not os.path.isfile(arial_path):
            return False, FileNotFoundError(f"Font file not found: {arial_path}")
        if not os.path.isfile(arial_bold_path):
            return False, FileNotFoundError(f"Bold font file not found: {arial_bold_path}")

        pdf.add_font('Arial', '', arial_path, uni=True)
        pdf.add_font('Arial-Bold', '', arial_bold_path, uni=True)
        pdf.set_font('Arial-Bold', '', 12)

        pdf.add_font('Arial', '', arial_path, uni=True)
        pdf.add_font('Arial-Bold', '', arial_bold_path, uni=True)
        pdf.set_font('Arial-Bold', '', 12)

    if sys.platform.startswith('lin'):
        current_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
        font_dir = os.path.join(current_dir, 'fonts')
        font_path = os.path.join(font_dir, 'DejaVuSans.ttf')
        bold_font_path = os.path.join(font_dir, 'DejaVuSansBold.ttf')

        # Debugging statements to ensure paths are correct
        print(f"Current Directory: {current_dir}")
        print(f"Font Directory: {font_dir}")
        print(f"Regular Font Path: {font_path}")
        print(f"Bold Font Path: {bold_font_path}")

        # Check if font files exist
        if not os.path.isfile(font_path):
            return False, FileNotFoundError(f"Font file not found: {font_path}")

        if not os.path.isfile(bold_font_path):
            return False, FileNotFoundError(f"Bold font file not found: {bold_font_path}")

        pdf.add_font('DejaVu', '', font_path, uni=True)
        pdf.add_font('DejaVu-Bold', '', bold_font_path, uni=True)
        pdf.set_font('DejaVu-Bold', '', 12)

    pdf.cell(200, 10, txt="Data from Textboxes", ln=True, align='C')

    for i, item in enumerate(data_prt):
        text = f"{item.get('idx', '')} | {item.get('name', '')} | {textboxes[i].text()}"
        if sys.platform.startswith('win32'):
            pdf.set_font('Arial', '', 12)

        if sys.platform.startswith('lin'):
            pdf.set_font('DejaVu', '', 12)

        pdf.multi_cell(0, 10, txt=text)

    pdf.add_page()
    if sys.platform.startswith('win32'):
        pdf.set_font('Arial', '', 12)

    if sys.platform.startswith('lin'):
        pdf.set_font('DejaVu', '', 12)
    pdf.cell(200, 10, txt="Result Data", ln=True, align='C')

    for row in result_data:
        text = " | ".join(row)
        if sys.platform.startswith('win32'):
            pdf.set_font('Arial', '', 12)

        if sys.platform.startswith('lin'):
            pdf.set_font('DejaVu', '', 12)
        pdf.multi_cell(0, 10, txt=text)

    try:
        output_dir = os.path.join(path, 'export')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, 'Untitled.pdf')
        pdf.output(output_path)
    except Exception as e:
        return False, e
    else:
        return True, None


def append_to_pdf(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to pdf file
    :param path: path to save to file
    :param data: Data to save to pdf file
    :param textboxes: Text boxes to save to pdf file
    :param result_data: Result data to save to pdf file
    :return: True on success, False on failure + Exception
    """
    try:
        from fpdf import FPDF
    except ImportError:
        return False, Exception('FPDF is not installed')

    try:
        if not data:
            return False, Exception('Data cannot be empty')

        data_prt = data[0].get('mdth', None)
        if not data_prt:
            return False, Exception('Data cannot be empty')

        if not textboxes:
            return False, Exception('Text boxes cannot be empty')

        if not os.path.exists(f'{path}/export/Untitled.pdf'):
            return save_to_pdf(path, data, textboxes, result_data)

        class PDF(FPDF):
            def header(self) -> None:
                self.set_font('Arial', 'B', 12)
                self.cell(0, 10, 'Data from Textboxes', 0, 1, 'C')

        # Create new PDF or read existing one
        pdf_path = os.path.join(path, 'export', 'Untitled.pdf')

        pdf = PDF()
        pdf.add_page()

        if sys.platform.startswith('win32'):
            arial_path = 'C:\\Windows\\Fonts\\arial.ttf'
            arial_bold_path = 'C:\\Windows\\Fonts\\arialbd.ttf'

            if not os.path.isfile(arial_path):
                return False, FileNotFoundError(f"Font file not found: {arial_path}")
            if not os.path.isfile(arial_bold_path):
                return False, FileNotFoundError(f"Bold font file not found: {arial_bold_path}")

            pdf.add_font('Arial', '', arial_path, uni=True)
            pdf.add_font('Arial-Bold', '', arial_bold_path, uni=True)
            pdf.set_font('Arial-Bold', '', 12)

        elif sys.platform.startswith('lin'):
            current_dir = os.path.abspath(os.path.dirname(__file__))
            font_dir = os.path.join(current_dir, 'fonts')
            font_path = os.path.join(font_dir, 'DejaVuSans.ttf')
            bold_font_path = os.path.join(font_dir, 'DejaVuSansBold.ttf')

            if not os.path.isfile(font_path):
                return False, FileNotFoundError(f"Font file not found: {font_path}")

            if not os.path.isfile(bold_font_path):
                return False, FileNotFoundError(f"Bold font file not found: {bold_font_path}")

            pdf.add_font('DejaVu', '', font_path, uni=True)
            pdf.add_font('DejaVu-Bold', '', bold_font_path, uni=True)
            pdf.set_font('DejaVu-Bold', '', 12)

        pdf.cell(200, 10, txt="Data from Textboxes", ln=True, align='C')

        for i, item in enumerate(data_prt):
            text = f"{item.get('idx', '')} | {item.get('name', '')} | {textboxes[i].text()}"
            if sys.platform.startswith('win32'):
                pdf.set_font('Arial', '', 12)
            elif sys.platform.startswith('lin'):
                pdf.set_font('DejaVu', '', 12)

            pdf.multi_cell(0, 10, txt=text)

        pdf.add_page()
        if sys.platform.startswith('win32'):
            pdf.set_font('Arial', '', 12)
        elif sys.platform.startswith('lin'):
            pdf.set_font('DejaVu', '', 12)
        pdf.cell(200, 10, txt="Result Data", ln=True, align='C')

        for row in result_data:
            text = " | ".join(row)
            if sys.platform.startswith('win32'):
                pdf.set_font('Arial', '', 12)
            elif sys.platform.startswith('lin'):
                pdf.set_font('DejaVu', '', 12)
            pdf.multi_cell(0, 10, txt=text)

        try:
            if os.path.exists(pdf_path):
                reader = PdfReader(pdf_path)
                writer = PdfWriter()

                for page in reader.pages:
                    writer.add_page(page)

                temp_pdf_path = os.path.join(path, 'export', 'Temp.pdf')
                pdf.output(temp_pdf_path)

                temp_reader = PdfReader(temp_pdf_path)
                for page in temp_reader.pages:
                    writer.add_page(page)

                with open(pdf_path, 'wb') as f:
                    writer.write(f)

                os.remove(temp_pdf_path)
            else:
                pdf.output(pdf_path)
        except Exception as e:
            return False, e
        else:
            return True, None
    except Exception as e:
        return False, e


def save_to_csv(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to csv file with csv format
    :param path: path to save to file
    :param data: Data to save to csv file with csv format
    :param textboxes: Text boxes to save to csv file with csv format
    :param result_data: Result data to save to csv file
    :return: True on success, False on failure + Exception
    """
    try:
        import csv
    except ImportError:
        return False, Exception('CSV module is not installed')

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    file = None
    try:
        file = open(f"{path}/export/Untitled.csv", mode='w', newline='', encoding='utf-8')
        writer = csv.writer(file)

        writer.writerow(['№ п/п', 'Показатель', 'Ответ субъекта'])
        for i, item in enumerate(data_prt):
            writer.writerow([item.get("idx", ""), item.get("name", ""), textboxes_ptr[i].text()])

        writer.writerow([])

        writer.writerow(['Result Data'])
        for row in result_data:
            writer.writerow(row)

    except Exception as e:
        return False, Exception(e)
    else:
        return True, None
    finally:
        if file:
            file.close()


def append_to_csv(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to csv file with csv format
    :param path: path to save to file
    :param data: Data to save to csv file with csv format
    :param textboxes: Text boxes to save to csv file with csv format
    :param result_data: Result data to save to csv file
    :return: True on success, False on failure + Exception
    """
    try:
        import csv
    except ImportError:
        return False, Exception('CSV module is not installed')

    try:
        if not os.path.exists(f'{path}/export/Untitled.csv'):
            return save_to_csv(path, data, textboxes, result_data)

        if data == []:
            return False, Exception('Data cannot be empty')

        data_prt = data[0].get('mdth', None)
        if not data_prt:
            return False, Exception('Data cannot be empty')

        textboxes_ptr: List = textboxes
        if not textboxes_ptr:
            return False, Exception('Text boxes cannot be empty')

        output_dir = os.path.join(path, 'export')
        os.makedirs(output_dir, exist_ok=True)
        file_path = os.path.join(output_dir, 'Untitled.csv')

        # Determine if file already exists
        file_exists = os.path.isfile(file_path)

        with open(file_path, mode='a', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)

            if not file_exists:
                # Write headers if file does not exist
                writer.writerow(['№ п/п', 'Показатель', 'Ответ субъекта'])

            # Write data from data_prt and textboxes
            for i, item in enumerate(data_prt):
                writer.writerow([item.get("idx", ""), item.get("name", ""), textboxes_ptr[i].text()])

            writer.writerow([])  # Add a blank row for separation

            # Write result data
            writer.writerow(['Result Data'])
            for row in result_data:
                writer.writerow(row)

        return True, None
    except Exception as e:
        return False, Exception(e)


def save_to_txt(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to txt file
    :param path: path to save to file
    :param data: Data to save to txt file
    :param textboxes: Text boxes to save to txt file
    :param result_data: Result data to save to txt file
    :return: True on success, False on failure + Exception
    """
    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    file = None
    try:
        file = open(f'{path}/export/Untitled.txt', mode='w', encoding='utf-8')
    except Exception as e:
        return False, Exception(e)
    else:
        for i, item in enumerate(data_prt):
            line = (f"№ п/п: {item.get('idx', '')}, "
                    f"Показатель: {item.get('name', '')}, "
                    f"Ответ субъекта: {textboxes[i].text()}\n")
            try:
                file.write(line)
            except Exception as e:
                return False, Exception(e)

        file.write("\n")
        if result_data:
            file.write("Result Data:\n")

            for row in result_data:
                line = " | ".join(row) + "\n"

                try:
                    file.write(line)
                except Exception as e:
                    return False, Exception(e)

        else:
            return True, None
    finally:
        if file:
            file.close()


def append_to_txt(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to txt file
    :param path: path to save to file
    :param data: Data to save to txt file
    :param textboxes: Text boxes to save to txt file
    :param result_data: Result data to save to txt file
    :return: True on success, False on failure + Exception
    """
    if not os.path.exists(f'{path}/export/Untitled.txt'):
        return save_to_txt(path, data, textboxes, result_data)

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    try:
        output_dir = os.path.join(path, 'export')
        os.makedirs(output_dir, exist_ok=True)
        file_path = os.path.join(output_dir, 'Untitled.txt')


        file_exists = os.path.isfile(file_path)

        with open(file_path, mode='a', encoding='utf-8') as file:
            if not file_exists:
                file.write("Data from Textboxes:\n")

            for i, item in enumerate(data_prt):
                line = (f"№ п/п: {item.get('idx', '')}, "
                        f"Показатель: {item.get('name', '')}, "
                        f"Ответ субъекта: {textboxes[i].text()}\n")
                file.write(line)

            file.write("\n")
            if result_data:
                file.write("Result Data:\n")
                for row in result_data:
                    line = " | ".join(row) + "\n"
                    file.write(line)

        return True, None
    except Exception as e:
        return False, Exception(e)


def save_to_xml(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data and result_data to xml file
    :param path: path to save to file
    :param data: Data to save to xml file
    :param textboxes: Text boxes to save to xml file
    :param result_data: Result data to save to xml file
    :return: True on success, False on failure + Exception
    """
    try:
        from xml.etree.ElementTree import Element, SubElement, ElementTree
    except ImportError:
        return False, Exception('XML is not installed')

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    root = Element('root')

    data_element = SubElement(root, 'data')
    for i, item in enumerate(data_prt):
        entry = SubElement(data_element, 'entry')
        SubElement(entry, 'number').text = item.get('idx', '')
        SubElement(entry, 'indicator').text = item.get('name', '')
        SubElement(entry, 'answer').text = textboxes_ptr[i].text()

    if result_data:
        result_element = SubElement(root, 'result_data')
        for row in result_data:
            entry = SubElement(result_element, 'entry')
            SubElement(entry, 'field1').text = row[0] if len(row) > 0 else ''
            SubElement(entry, 'field2').text = row[1] if len(row) > 1 else ''
            SubElement(entry, 'field3').text = row[2] if len(row) > 2 else ''

    tree = ElementTree(root)
    try:
        tree.write(f'{path}/export/Untitled.xml', encoding='utf-8', xml_declaration=True)
    except Exception as e:
        return False, Exception(e)
    else:
        return True, None


def append_to_xml(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data and result_data to xml file
    :param path: path to save to file
    :param data: Data to save to xml file
    :param textboxes: Text boxes to save to xml file
    :param result_data: Result data to save to xml file
    :return: True on success, False on failure + Exception
    """
    try:
        import xml.etree.ElementTree as ET
    except ImportError:
        return False, Exception('XML is not installed')

    if not os.path.exists(f'{path}/export/Untitled.xml'):
        return save_to_xml(path, data, textboxes, result_data)

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    output_dir = os.path.join(path, 'export')
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, 'Untitled.xml')

    try:
        if os.path.exists(file_path):
            tree = ET.parse(file_path)
            root = tree.getroot()
        else:
            root = ET.Element('root')

        data_element = root.find('data')
        if data_element is None:
            data_element = ET.SubElement(root, 'data')

        for i, item in enumerate(data_prt):
            entry = ET.SubElement(data_element, 'entry')
            ET.SubElement(entry, 'number').text = item.get('idx', '')
            ET.SubElement(entry, 'indicator').text = item.get('name', '')
            ET.SubElement(entry, 'answer').text = textboxes_ptr[i].text()

        if result_data:
            result_element = root.find('result_data')
            if result_element is None:
                result_element = ET.SubElement(root, 'result_data')
            for row in result_data:
                entry = ET.SubElement(result_element, 'entry')
                ET.SubElement(entry, 'field1').text = row[0] if len(row) > 0 else ''
                ET.SubElement(entry, 'field2').text = row[1] if len(row) > 1 else ''
                ET.SubElement(entry, 'field3').text = row[2] if len(row) > 2 else ''

        tree = ET.ElementTree(root)
        tree.write(file_path, encoding='utf-8', xml_declaration=True)
        return True, None
    except Exception as e:
        return False, Exception(e)


def save_to_json(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to json file with json format
    :param path: path to save to file
    :param data: Data to save to json file with json format
    :param textboxes: Text boxes to save to json file with json format
    :param result_data: Result data to save to json file
    :return: True on success, False on failure + Exception
    """
    try:
        import json
    except ImportError:
        return False, Exception('JSON is not installed')

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    # Prepare data
    data_list = [{'№ п/п': item.get("idx", ""), 'Показатель': item.get("name", ""),
                  'data': textboxes_ptr[i].text()} for i, item in enumerate(data_prt)]

    result_list = [{'Field1': row[0], 'Field2': row[1], 'Field3': row[2]} for row in result_data] if result_data else []

    combined_data = {'data': data_list, 'result_data': result_list}

    file = None
    try:
        file = open(f'{path}/export/Untitled.json', mode='w', encoding='utf-8')
    except Exception as e:
        return False, Exception(e)
    else:
        try:
            json.dump(combined_data, file, indent=4, ensure_ascii=False)
        except Exception as e:
            return False, Exception(e)
        else:
            return True, None
    finally:
        if file:
            file.close()


def append_to_json(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to json file with json format
    :param path: path to save to file
    :param data: Data to save to json file with json format
    :param textboxes: Text boxes to save to json file with json format
    :param result_data: Result data to save to json file
    :return: True on success, False on failure + Exception
    """
    try:
        import json
    except ImportError:
        return False, Exception('JSON is not installed')

    if not os.path.exists(f'{path}/export/Untitled.json'):
        return save_to_json(path, data, textboxes, result_data)

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    output_dir = os.path.join(path, 'export')
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, 'Untitled.json')

    try:
        # Load existing JSON data if file exists
        existing_data = {}
        if os.path.exists(file_path):
            with open(file_path, mode='r', encoding='utf-8') as f:
                existing_data = json.load(f)

        # Prepare new data to add
        data_list = [{'№ п/п': item.get("idx", ""), 'Показатель': item.get("name", ""),
                      'data': textboxes_ptr[i].text()} for i, item in enumerate(data_prt)]

        result_list = [{'Field1': row[0], 'Field2': row[1], 'Field3': row[2]} for row in
                       result_data] if result_data else []

        combined_data = {'data': data_list, 'result_data': result_list}

        # Merge new data with existing data
        existing_data.update(combined_data)

        with open(file_path, mode='w', encoding='utf-8') as file:
            json.dump(existing_data, file, indent=4, ensure_ascii=False)

        return True, None
    except Exception as e:
        return False, Exception(e)


def save_to_html(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to html file with html format with html format
    :param path: path to save to file
    :param data: Data to save to html file with html format
    :param textboxes: Text boxes to save to html file with html format
    :param result_data: Result data to save to html file
    :return: True on success, False on failure + Exception
    """
    try:
        import pandas as pd
    except ImportError:
        return False, Exception('Pandas is not installed')

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    # Prepare data
    data_list = [{'№ п/п': item.get("idx", ""), 'Показатель': item.get("name", ""),
                  'Ответ субъекта': textboxes_ptr[i].text()} for i, item in enumerate(data_prt)]
    data_df = pd.DataFrame(data_list)
    data_html = data_df.to_html(index=False)
    result_df = pd.DataFrame(result_data) if result_data else pd.DataFrame()
    result_html = result_df.to_html(index=False)
    try:
        with open(f'{path}/export/Untitled.html', mode='w', encoding='utf-8') as file:
            file.write("<html><head><title>Output HTML</title></head><body>\n")
            file.write("<h2>Data from Textboxes</h2>\n")
            file.write(data_html)
            file.write("<h2>Result Data</h2>\n")
            file.write(result_html)
            file.write("</body></html>")
    except Exception as e:
        return False, Exception(e)
    else:
        return True, None


def append_to_html(path: str, data: List, textboxes: List, result_data: List[List[str]]) -> [bool, Exception]:
    """
    Save data to html file with html format
    :param path: path to save to file
    :param data: Data to save to html file with html format
    :param textboxes: Text boxes to save to html file with html format
    :param result_data: Result data to save to html file
    :return: True on success, False on failure + Exception
    """
    try:
        import pandas as pd
    except ImportError:
        return False, Exception('Pandas is not installed')

    if not os.path.exists(f'{path}/export/Untitled.html'):
        return save_to_html(path, data, textboxes, result_data)

    if data == []:
        return False, Exception('Data cannot be empty')

    data_prt = data[0].get('mdth', None)
    if not data_prt:
        return False, Exception('Data cannot be empty')

    textboxes_ptr: List = textboxes
    if not textboxes_ptr:
        return False, Exception('Text boxes cannot be empty')

    output_dir = os.path.join(path, 'export')
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, 'Untitled.html')

    # Prepare data
    data_list = [{'№ п/п': item.get("idx", ""), 'Показатель': item.get("name", ""),
                  'Ответ субъекта': textboxes_ptr[i].text()} for i, item in enumerate(data_prt)]
    data_df = pd.DataFrame(data_list)
    data_html = data_df.to_html(index=False)

    result_df = pd.DataFrame(result_data) if result_data else pd.DataFrame()
    result_html = result_df.to_html(index=False)

    try:
        with open(file_path, mode='a', encoding='utf-8') as file:
            if not os.path.exists(file_path):
                file.write("<html><head><title>Output HTML</title></head><body>\n")

            file.write("<h2>Data from Textboxes</h2>\n")
            file.write(data_html)
            file.write("<h2>Result Data</h2>\n")
            file.write(result_html)
            file.write("</body></html>")
        return True, None
    except Exception as e:
        return False, Exception(e)


def generate_mdth_file(data, filename: str, mode: str = "w", encoding: str = "utf-8") -> [bool, Exception]:
    """
    Generate markdown file from data
    :param data: data to save
    :param filename: file name to save markdown
    :param mode: (Optional) mode to save file
    :param encoding: (Optional) encoding to save file
    :return: True or False with Exception
    """
    try:
        import json
    except ImportError:
        return False, Exception('JSON is not installed')
    try:
        json_data = json.dumps(data, ensure_ascii=False, indent=4)
    except Exception as e:
        return False, e

    file = None
    try:
        file = open(filename, mode, encoding=encoding)
    except Exception as e:
        return False, Exception(e)
    else:
        file.write(json_data)
        return True, None
    finally:
        if file:
            file.close()


def load_mdth_file(filename: str) -> [bool, Exception]:
    """
    Load markdown file from file
    :param filename:  file name to load
    :return: load markdown
    """
    import json
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            content = f.read()
            data = json.loads(content)
            return data

    except FileNotFoundError:
        print(f"Файл '{filename}' не найден.")
        return None
    except json.JSONDecodeError as e:
        print(f"Ошибка при чтении файла '{filename}': {e}")
        return None


def create_project(dir_path: str, file_type: str = "*") -> [bool, Exception]:
    """
    Create project directory
    :param dir_path: project directory
    :param file_type: (Optional) file type
    :return: True on success, False on failure + Exception
    """
    try:
        import json
    except ImportError:
        return False, Exception('JSON is not installed')

    if not os.path.exists(dir_path):
        os.makedirs(dir_path)

    project_data = {
        "directory": dir_path,
        "file_type": file_type
    }

    try:
        projmd_file = os.path.join(dir_path, "proj.projmd")
        with open(projmd_file, 'w') as f:
            json.dump(project_data, f, indent=4)
    except Exception as e:
        return False, Exception(e)

    from utils.configuration_config_mdt import ConfigurationMDTH
    try:
        ConfigurationMDTH.create_configuration_config_mdt(dir_path).save_as_json()
    except Exception as e:
        return False, Exception(e)
    else:
        return True, None


def open_project(dir_path: str, file_type: str = "*") -> [bool, Exception]:
    """
    Open project directory
    :param dir_path: project directory
    :param file_type: (Optional) file type
    :return: True on success, False on failure + Exception
    """
    try:
        import json
    except ImportError:
        return False, Exception('JSON is not installed')

    projmd_file = os.path.join(dir_path, "proj.projmd")
    config_file = os.path.join(dir_path, "configuration_config_mdt.json")

    if not os.path.exists(config_file):
        return False, Exception('Configuration file does not exist')

    if not os.path.exists(projmd_file):
        return False, Exception('Project file does not exist')

    try:
        with open(projmd_file, 'r') as f:
            project_data = json.load(f)
    except Exception as e:
        return False, Exception(e)

    return project_data, None


def check_main_dirs(path: str) -> None:
    """
    Check if the main directories exist
    :param path: path to the main directories
    :return: None
    """
    import os
    if not os.path.exists(os.path.join(path, 'export')):
        os.mkdir(os.path.join(path, 'export'))


def copy_file(file_path: str, dir_path: str = None) -> [bool, Exception]:
    """
    Copy file from source to destination directory
    :param file_path: file name to copy to destination directory
    :param dir_path: destination directory (Optional)
    :return: True on success, False on failure + Exceptions
    """

    try:
        import shutil
    except ImportError:
        return False, Exception("Module 'shutil' is not installed")

    if os.path.isfile(file_path):
        if not dir_path:
            dir_path: str = os.path.dirname(file_path)
        base_name: str = os.path.basename(file_path)
        copy_path: str = os.path.join(dir_path, f"tmp_{base_name}")
        try:
            shutil.copy(file_path, copy_path)
        except Exception as e:
            return False, Exception(e)
        else:
            return True, None
