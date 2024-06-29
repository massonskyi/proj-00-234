import os

import pandas as pd
import pdfplumber
from PySide6 import (
    QtWidgets,
    QtCore,
)
from PySide6.QtCore import (
    Signal,
    QModelIndex
)
from PySide6.QtGui import QAction
from docx import Document

from utils.s2f import load_mdth_file


class CustomFileManager(QtWidgets.QWidget):
    """
    Custom file manager widget.
    """
    file_selected = Signal(list)  # Сигнал для выбора файла
    filepath_selected = Signal(str)  # Сигнал для выбора пути файла

    def __init__(self, path: str = None, parent=None) -> None:
        """
        Initialize the CustomFileManager widget.
        """
        super().__init__(parent)
        self.setWindowTitle('File Manager')

        layout: QtWidgets.QVBoxLayout = QtWidgets.QVBoxLayout(self)

        self.model: QtWidgets.QFileSystemModel = QtWidgets.QFileSystemModel()
        self.model.setRootPath(QtCore.QDir.rootPath())

        self.tree: QtWidgets.QTreeView = QtWidgets.QTreeView()
        self.tree.setModel(self.model)
        if path:
            self.tree.setRootIndex(self.model.index(path))
        self.tree.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)

        self.tree.setColumnWidth(0, 250)
        self.tree.setColumnWidth(1, 100)
        self.tree.setColumnWidth(2, 100)
        self.tree.setColumnWidth(3, 150)

        self.tree.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self.open_menu)
        self.tree.doubleClicked.connect(self.select_file)

        self.create_file_action: QAction = QAction("Create File", self)
        self.create_dir_action: QAction = QAction("Create Directory", self)
        self.delete_action: QAction = QAction("Delete", self)
        self.rename_action: QAction = QAction("Rename", self)

        self.create_file_action.triggered.connect(self.create_file)
        self.create_dir_action.triggered.connect(self.create_directory)
        self.delete_action.triggered.connect(self.delete_item)
        self.rename_action.triggered.connect(self.rename_item)

        layout.addWidget(self.tree)
        button_layout: QtWidgets.QHBoxLayout = QtWidgets.QHBoxLayout()
        layout.addLayout(button_layout)

    def load_file(self, file_path: str) -> None:
        """
        Load the file to the widget.
        :param file_path: path of the file to be loaded
        :return: None
        """
        if file_path:
            self.filepath_selected.emit(file_path)
            data = []
            file_ext: str = file_path.lower()

            if file_ext.endswith('.mdth'):
                data = [{"mdth": load_mdth_file(file_path)}]
            elif file_ext.endswith(('.png', '.jpg', '.bmp')):
                data = [{'image_path': file_path}]
            elif file_ext.endswith('.csv'):
                data = pd.read_csv(file_path).to_dict(orient='records')
            elif file_ext.endswith('.xlsx') or file_ext.endswith('.xls'):
                data = pd.read_excel(file_path).to_dict(orient='records')
            elif file_ext.endswith('.docx'):
                doc: Document = Document(file_path)
                data_temp: list = []
                for table in doc.tables:
                    table_data: list = []
                    for row in table.rows:
                        row_data: list = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        table_data.append(row_data)
                    data_temp.append(table_data)
                data = [{"dataword": data_temp}]

            elif file_ext.endswith('.pdf'):
                with pdfplumber.open(file_path) as pdf:
                    tmp: list = []
                    for page in pdf.pages:
                        tmp.append(page.extract_text())

                    data = [{"pdf": tmp}]
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data.append(f.read())

            self.file_selected.emit(data)

    def select_file(self, index: QModelIndex) -> None:
        """
        Select the file to be open.
        :param index: index of the selected file
        :param path: path of the selected file
        :return: None
        """
        file_path = self.model.filePath(index)

        if not os.path.isfile(file_path):
            return

        if file_path:
            self.filepath_selected.emit(file_path)
            data = []
            file_ext: str = file_path.lower()

            if file_ext.endswith('.mdth'):
                data = [{"mdth": load_mdth_file(file_path)}]
            elif file_ext.endswith(('.png', '.jpg', '.bmp')):
                data = [{'image_path': file_path}]
            elif file_ext.endswith('.csv'):
                data = pd.read_csv(file_path).to_dict(orient='records')
            elif file_ext.endswith('.xlsx') or file_ext.endswith('.xls'):
                data = pd.read_excel(file_path).to_dict(orient='records')
            elif file_ext.endswith('.docx'):
                doc: Document = Document(file_path)
                data_temp: list = []
                for table in doc.tables:
                    table_data: list = []
                    for row in table.rows:
                        row_data: list = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        table_data.append(row_data)
                    data_temp.append(table_data)
                data = [{"dataword": data_temp}]

            elif file_ext.endswith('.pdf'):
                with pdfplumber.open(file_path) as pdf:
                    tmp: list = []
                    for page in pdf.pages:
                        tmp.append(page.extract_text())

                    data = [{"pdf": tmp}]
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data.append(f.read())

            self.file_selected.emit(data)

    def open_menu(self, position: QtCore.QPoint) -> None:
        """
        Open the file menu.
        :param position: position of the menu
        :return: None
        """
        indexes: QModelIndex = self.tree.selectedIndexes()
        if indexes:
            menu: QtWidgets.QMenu = QtWidgets.QMenu()
            menu.addAction(self.create_file_action)
            menu.addAction(self.create_dir_action)
            menu.addAction(self.delete_action)
            menu.addAction(self.rename_action)
            menu.exec_(self.tree.viewport().mapToGlobal(position))

    def create_file(self) -> None:
        """
        Create the file manager.
        :return: None
        """
        index: QModelIndex = self.tree.currentIndex()
        if not index.isValid():
            return

        dir_path: str = self.model.filePath(index)
        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Create File", dir_path)
        if file_path:
            open(file_path, 'w').close()  # Создаем пустой файл

    def create_directory(self) -> None:
        """
        Create the directory manager.
        :return: None
        """
        index: QModelIndex = self.tree.currentIndex()
        if not index.isValid():
            return

        dir_path: str = self.model.filePath(index)
        dir_name, ok = QtWidgets.QInputDialog.getText(self, "Create Directory", "Directory Name:")
        if ok and dir_name:
            os.mkdir(os.path.join(dir_path, dir_name))

    def delete_item(self) -> None:
        """
        Delete the selected file.
        :return: None
        """
        index: QModelIndex = self.tree.currentIndex()
        if not index.isValid():
            return

        file_path: str = self.model.filePath(index)
        if os.path.isdir(file_path):
            os.rmdir(file_path)
        else:
            os.remove(file_path)

    def rename_item(self) -> None:
        """
        Rename the selected file.
        :return: None
        """
        index: QModelIndex = self.tree.currentIndex()
        if not index.isValid():
            return

        file_path: str = self.model.filePath(index)
        new_name, ok = QtWidgets.QInputDialog.getText(self, "Rename Item", "New Name:")
        if ok and new_name:
            new_path = os.path.join(os.path.dirname(file_path), new_name)
            os.rename(file_path, new_path)
