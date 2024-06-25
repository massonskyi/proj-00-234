import csv
import json

import pandas as pd
from PySide6.QtCore import Qt, QEvent
from PySide6.QtGui import QIntValidator, QPixmap
from PySide6.QtWidgets import QWidget, QVBoxLayout, QScrollArea, QLabel, QFrame, QPushButton, QGridLayout, QLineEdit, \
    QTextEdit, QHBoxLayout, QMenu, QSizePolicy, QSpacerItem
from docx import Document
from fpdf import FPDF

from utils.loaders import load_icon


class ScrollableContainer(QWidget):
    def __init__(self, data: list = None, file_type: str = None, parent=None):
        super().__init__(parent)

        self.icons = {
            'save_word': load_icon('./assets/save2/sword.png'),
            'save_excel': load_icon('./assets/save2/sexcel.png'),
            'save_pdf': load_icon('./assets/save2/spdf.png'),
            'save_csv': load_icon('./assets/save2/scsv.png'),
            'save_json': load_icon('./assets/save2/sjson.png'),
            'save_html': load_icon('./assets/save2/shtml.png'),
            'save_txt': load_icon('./assets/save2/stxt.png'),
            'save_xml': load_icon('./assets/save2/sxml.png'),
        }
        self.buttons_name = {
            'save_word': 'Сохранить как Word',
            'save_excel': 'Сохранить как Excel',
            'save_pdf': 'Сохранить как PDF',
            'save_csv': 'Сохранить как CSV',
            'save_json': 'Сохранить как JSON',
            'save_html': 'Сохранить как HTML',
            'save_txt': 'Сохранить как TXT',
            'save_xml': 'Сохранить как XML',
        }
        self.textboxes = None
        self.data = data if data else []

        main_layout = QVBoxLayout(self)

        self.scroll_area = QScrollArea(self)
        self.scroll_area.setWidgetResizable(True)

        self.scroll_content = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_content)

        if file_type == 'text':
            self.load_text_content(self.scroll_layout)
        elif file_type == 'image':
            self.load_image_content(self.scroll_layout)
        else:
            self.load_default_content(self.scroll_layout)

        self.scroll_area.setWidget(self.scroll_content)

        main_layout.addWidget(self.scroll_area)

        # Add additional container with a button
        self.additional_container = QFrame(self)
        self.additional_container.setFrameShape(QFrame.StyledPanel)
        self.additional_layout = QVBoxLayout(self.additional_container)

        self.additional_buttons_container = QFrame(self)
        self.additional_buttons_layout = QHBoxLayout(self.additional_buttons_container)
        self.additional_buttons_layout.setContentsMargins(0, 0, 0, 0)  # Ensure no extra margins

        self.add_initial_buttons()

        self.additional_layout.addWidget(self.additional_buttons_container)
        self.additional_container.hide()

        main_layout.addWidget(self.additional_container)

        # Connect signals for scrolling
        self.scroll_area.verticalScrollBar().valueChanged.connect(self.check_scroll_position)

    def add_initial_buttons(self):
        buttons = []
        for it, (key, icon) in enumerate(self.icons.items()):
            button = QPushButton(icon, self.buttons_name.get(key))
            button.setObjectName(key)
            button.setFixedSize(175, 40)
            button.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
            button.clicked.connect(self.save)
            buttons.append(button)

            if len(buttons) <= 3:
                self.additional_buttons_layout.addWidget(button)

        if len(buttons) > 3:
            menu_button = QPushButton("...", self.additional_buttons_container)
            menu_button.setFixedSize(175, 40)
            menu_button.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
            self.additional_buttons_layout.addWidget(menu_button)

            menu = QMenu(self)
            for button in buttons[3:]:
                action = menu.addAction(button.icon(),self.buttons_name.get(button.objectName()))
                action.setObjectName(button.objectName())
                action.triggered.connect(self.save)  # Connect the QAction's triggered signal to button's click slot

            menu_button.setMenu(menu)

        # Add spacers to ensure even distribution of buttons
        spacer_left = QSpacerItem(20, 40, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        spacer_right = QSpacerItem(20, 40, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.additional_buttons_layout.insertItem(0, spacer_left)
        self.additional_buttons_layout.addItem(spacer_right)

    def save(self):
        sender = self.sender()
        if sender.objectName() == 'save_word':
            self.save_to_word()
        elif sender.objectName() == 'save_excel':
            self.save_to_excel()
        elif sender.objectName() == 'save_pdf':
            self.save_to_pdf()
        elif sender.objectName() == 'save_csv':
            self.save_to_csv()
        elif sender.objectName() == 'save_json':
            self.save_to_json()
        elif sender.objectName() == 'save_html':
            self.save_to_html()
        elif sender.objectName() == 'save_txt':
            print(f"Button {sender.objectName()} clicked")
    def save_to_word(self):
        document = Document()
        document.add_heading('Data from Textboxes', level=1)

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№ п/п'
        hdr_cells[1].text = 'Показатель'
        hdr_cells[2].text = 'Ответ субъекта'

        for i, item in enumerate(self.data[0].get('mdth')):
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("№ п/п", "")
            row_cells[1].text = item.get("Показатель", "")
            row_cells[2].text = self.textboxes[i].text()

        document.save('output.docx')
        print("Document saved as output.docx")

    def save_to_excel(self):
        data = [{'№ п/п': item.get("№ п/п", ""), 'Показатель': item.get("Показатель", ""),
                 'Ответ субъекта': self.textboxes[i].text()} for i, item in enumerate(self.data[0].get('mdth'))]
        df = pd.DataFrame(data)
        df.to_excel('output.xlsx', index=False)
        print("Document saved as output.xlsx")

    def save_to_pdf(self):
        from fpdf import FPDF

        class PDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 12)
                self.cell(0, 10, 'Data from Textboxes', 0, 1, 'C')

        pdf = PDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Data from Textboxes", ln=True, align='C')

        for i, item in enumerate(self.data[0].get('mdth')):
            text = f"{item.get('№ п/п', '')} | {item.get('Показатель', '')} | {self.textboxes[i].text()}"
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, txt=text.encode('latin-1', 'replace').decode('latin-1'))

        pdf.output("output.pdf")
        print("Document saved as output.pdf")

    def save_to_csv(self):
        with open('output.csv', mode='w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            writer.writerow(['№ п/п', 'Показатель', 'Ответ субъекта'])
            for i, item in enumerate(self.data[0].get('mdth')):
                writer.writerow([item.get("№ п/п", ""), item.get("Показатель", ""), self.textboxes[i].text()])
        print("Document saved as output.csv")

    def save_to_json(self):
        data = [{'№ п/п': item.get("№ п/п", ""), 'Показатель': item.get("Показатель", ""),
                 'Ответ субъекта': self.textboxes[i].text()} for i, item in enumerate(self.data[0].get('mdth'))]
        with open('output.json', 'w', encoding='utf-8') as file:
            json.dump(data, file, indent=4, ensure_ascii=False)
        print("Document saved as output.json")

    def save_to_html(self):
        data = [{'№ п/п': item.get("№ п/п", ""), 'Показатель': item.get("Показатель", ""),
                 'Ответ субъекта': self.textboxes[i].text()} for i, item in enumerate(self.data[0].get('mdth'))]
        df = pd.DataFrame(data)
        df.to_html('output.html', index=False)
        print("Document saved as output.html")
    def check_scroll_position(self):
        scroll_bar = self.scroll_area.verticalScrollBar()
        max_scroll = scroll_bar.maximum()
        current_scroll = scroll_bar.value()
        if current_scroll >= max_scroll:
            self.additional_container.show()
        else:
            self.additional_container.hide()

    def load_text_content(self, layout):
        text_edit = QTextEdit()
        for item in self.data:
            text_edit.append(item)
        layout.addWidget(text_edit)

    def load_image_content(self, layout):
        for item in self.data:
            image_label = QLabel()
            pixmap = QPixmap(item.get('image_path', ''))
            if not pixmap.isNull():
                image_label.setPixmap(pixmap.scaledToWidth(400))  # Adjust width as needed
                layout.addWidget(image_label)

    def load_default_content(self, layout):
        grid_layout = QGridLayout()
        grid_layout.setColumnStretch(1, 1)

        int_validator = QIntValidator(0, 10, self)

        self.textboxes = []

        for row, item in enumerate(self.data[0].get('mdth')):
            number_label = QLabel(item.get("№ п/п", ""))
            indicator_label = QLabel(item.get("Показатель", ""))
            answer_textbox = QLineEdit()
            answer_textbox.setPlaceholderText(item.get("Ответ субъекта", ""))
            answer_textbox.setValidator(int_validator)
            answer_textbox.installEventFilter(self)

            self.textboxes.append(answer_textbox)

            number_label.setMinimumWidth(50)
            number_label.setAlignment(Qt.AlignCenter)

            indicator_label.setWordWrap(True)
            grid_layout.addWidget(number_label, row, 0)
            grid_layout.addWidget(indicator_label, row, 1)
            grid_layout.addWidget(answer_textbox, row, 2)

        layout.addLayout(grid_layout)

    def eventFilter(self, source, event):
        if not isinstance(event, QEvent):
            return

        if event.type() == QEvent.Type.KeyPress and isinstance(source, QLineEdit):
            current_index = self.textboxes.index(source)
            if event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter:
                next_index = (current_index + 1) % len(self.textboxes)
                self.textboxes[next_index].setFocus()
                self.scroll_area.ensureWidgetVisible(self.textboxes[next_index])
                return True

            elif event.key() == Qt.Key.Key_Up:
                next_index = (current_index - 1) % len(self.textboxes)
                self.textboxes[next_index].setFocus()
                self.scroll_area.ensureWidgetVisible(self.textboxes[next_index])
                return True

            elif event.key() == Qt.Key.Key_Down:
                next_index = (current_index + 1) % len(self.textboxes)
                self.textboxes[next_index].setFocus()
                self.scroll_area.ensureWidgetVisible(self.textboxes[next_index])
                return True

            elif event.key() == Qt.Key.Key_Up and event.modifiers() == Qt.ControlModifier:
                self.parentWidget().focusPreviousChild()
                return True

            elif event.key() == Qt.Key.Key_Down and event.modifiers() == Qt.ControlModifier:
                self.parentWidget().focusNextChild()
                return True
        return super().eventFilter(source, event)

    def update_content(self, data):
        self.data = data
        self.update_ui()

    def update_ui(self):
        for i in reversed(range(self.scroll_layout.count())):
            if isinstance(self.scroll_layout.itemAt(i), QGridLayout):
                for item in reversed(range(self.scroll_layout.itemAt(i).count())):
                    widget = self.scroll_layout.itemAt(i).itemAt(item).widget()
                    if widget:
                        widget.setParent(None)
            else:
                widget = self.scroll_layout.itemAt(i).widget()
                if widget:
                    widget.setParent(None)

        if self.data and isinstance(self.data, list):
            if self.data[0] and isinstance(self.data[0], dict):
                first_key = next(iter(self.data[0].keys()))
                if first_key == 'image_path':
                    self.load_image_content(self.scroll_layout)

                if first_key == 'mdth':
                    self.load_default_content(self.scroll_layout)
            else:
                self.load_text_content(self.scroll_layout)
